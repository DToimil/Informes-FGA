import ast
import os
import re

import docx
import streamlit as st
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ==========================================
# FUNCIONES AUXILIARES COMUNES
# ==========================================

SECCIONES_DR_RUTA = [
    "CÁMARA DE LLAMADAS",
    "SALIDAS",
    "CRONOMETRAJE TRANSP.",
    "CRONOMETRAJE MANUAL",
    "LLEGADAS",
    "CUENTAVUELTAS",
    "JUECES DE MARCHA",
    "JUECES DE RECORRIDO",
    "SECRET. COMPETICIÓN",
    "OTROS",
]


def eliminar_fila(row):
    """Elimina una fila limpiamente desde el XML de Word."""
    tr = row._tr
    if tr.getparent() is not None:
        tr.getparent().remove(tr)


def eliminar_tabla(tabla):
    """Elimina una tabla completa del documento."""
    tbl = tabla._tbl
    if tbl.getparent() is not None:
        tbl.getparent().remove(tbl)


def limpiar_nombre_archivo(nombre):
    nombre = str(nombre or "Competicion").strip()
    nombre = re.sub(r'[<>:"/\\|?*]+', "-", nombre)
    nombre = re.sub(r"\s+", " ", nombre).strip()
    return nombre


def obtener_valor(datos, clave, default=""):
    """Acepta claves con o sin {{...}} para facilitar correcciones manuales."""
    if clave in datos:
        return datos[clave]
    return datos.get(f"{{{{{clave}}}}}", default)


def tiene_contenido(valor):
    if isinstance(valor, bool):
        return valor
    if valor is None:
        return False
    if isinstance(valor, dict):
        return any(tiene_contenido(item) for item in valor.values())
    if isinstance(valor, (list, tuple, set)):
        return any(tiene_contenido(item) for item in valor)
    return bool(str(valor).strip())


def normalizar_marca(valor, estilo="x"):
    if isinstance(valor, bool):
        activo = valor
    else:
        texto = str(valor).strip().upper()
        activo = texto in {"1", "X", "SI", "SÍ", "TRUE", "VERDADERO", "YES"}

    if not activo:
        return ""
    return "x" if estilo == "x" else "SI"


def copiar_formato_run(run_origen, run_destino):
    """Copia solo lo esencial del formato del run original."""
    if run_origen is None:
        return

    run_destino.font.name = run_origen.font.name
    run_destino.font.size = run_origen.font.size
    run_destino.font.bold = run_origen.font.bold
    run_destino.font.italic = run_origen.font.italic
    run_destino.font.underline = run_origen.font.underline
    if run_origen.font.color and run_origen.font.color.rgb:
        run_destino.font.color.rgb = run_origen.font.color.rgb


def obtener_run_referencia(parrafo):
    for run in parrafo.runs:
        if run.text.strip():
            return run
    if parrafo.runs:
        return parrafo.runs[0]
    return None


def establecer_texto_manteniendo_formato(parrafo, texto):
    """Reescribe un párrafo manteniendo el estilo del primer run útil."""
    run_referencia = obtener_run_referencia(parrafo)
    parrafo.clear()
    nuevo_run = parrafo.add_run("" if texto is None else str(texto))
    copiar_formato_run(run_referencia, nuevo_run)


def establecer_texto_celda(celda, texto):
    if celda.paragraphs:
        establecer_texto_manteniendo_formato(celda.paragraphs[0], texto)
        for parrafo_extra in celda.paragraphs[1:]:
            establecer_texto_manteniendo_formato(parrafo_extra, "")
    else:
        celda.text = "" if texto is None else str(texto)


def reemplazar_manteniendo_formato_estricto(parrafo, datos):
    """Reemplaza datos manteniendo negritas, cursivas y color del párrafo."""
    if not parrafo.text or "{{" not in parrafo.text:
        return

    for run in parrafo.runs:
        for clave, valor in datos.items():
            if clave in run.text:
                run.text = run.text.replace(clave, str(valor))

        run.text = re.sub(r"\{\{.*?\}\}", "", run.text)

    if "{{" in parrafo.text:
        texto_completo = parrafo.text
        run_referencia = obtener_run_referencia(parrafo)

        for clave, valor in datos.items():
            if clave in texto_completo:
                texto_completo = texto_completo.replace(clave, str(valor))

        texto_completo = re.sub(r"\{\{.*?\}\}", "", texto_completo)

        parrafo.clear()
        nuevo_run = parrafo.add_run(texto_completo)
        copiar_formato_run(run_referencia, nuevo_run)


def activar_checkbox_por_posicion(doc, indice_real, activar=True):
    """Activa o desactiva los cuadraditos de Word por su orden de aparición."""
    checkboxes = doc.element.xpath(".//w:checkBox")
    if indice_real < len(checkboxes):
        cb = checkboxes[indice_real]
        valor = "1" if activar else "0"

        default_val = cb.find(qn("w:default"))
        if default_val is not None:
            default_val.set(qn("w:val"), valor)

        checked_val = cb.find(qn("w:checked"))
        if checked_val is None:
            checked_val = OxmlElement("w:checked")
            cb.append(checked_val)
        checked_val.set(qn("w:val"), valor)


def limpiar_tabla_por_etiquetas(tabla, datos, secciones_conocidas):
    filas_a_borrar = []
    cabecera_actual = None
    seccion_con_datos = False

    for fila in tabla.rows:
        textos_celdas = [celda.text.strip() for celda in fila.cells]
        texto_fila = "".join(textos_celdas)
        texto_celda_0 = textos_celdas[0].upper() if textos_celdas else ""

        if not texto_fila:
            continue

        if texto_celda_0 in secciones_conocidas:
            if cabecera_actual is not None and not seccion_con_datos:
                filas_a_borrar.append(cabecera_actual)
            cabecera_actual = fila
            seccion_con_datos = False
            continue

        if "{{" in texto_fila:
            celda_nombre = fila.cells[1].text if len(fila.cells) > 1 else fila.cells[0].text
            etiqueta_nombre = re.search(r"\{\{.*?_NOMBRE\}\}", celda_nombre)

            if etiqueta_nombre:
                etiqueta = etiqueta_nombre.group()
                if etiqueta not in datos or not str(datos[etiqueta]).strip():
                    filas_a_borrar.append(fila)
                    continue
                seccion_con_datos = True
            else:
                todas_etiquetas = re.findall(r"\{\{.*?\}\}", texto_fila)
                todas_vacias = True
                for etiqueta in todas_etiquetas:
                    if etiqueta in datos and str(datos[etiqueta]).strip():
                        todas_vacias = False
                        break
                if todas_vacias and todas_etiquetas:
                    filas_a_borrar.append(fila)
                    continue
                seccion_con_datos = True

        for celda in fila.cells:
            for parrafo in celda.paragraphs:
                reemplazar_manteniendo_formato_estricto(parrafo, datos)

    if cabecera_actual is not None and not seccion_con_datos:
        filas_a_borrar.append(cabecera_actual)

    for fila in filas_a_borrar:
        eliminar_fila(fila)

    filas_vacias_consecutivas = 0
    filas_rayas_a_borrar = []

    for fila in tabla.rows:
        textos_celdas = [celda.text.strip() for celda in fila.cells]
        if not "".join(textos_celdas):
            filas_vacias_consecutivas += 1
            if filas_vacias_consecutivas > 1:
                filas_rayas_a_borrar.append(fila)
        else:
            filas_vacias_consecutivas = 0

    if tabla.rows:
        textos_ultima_fila = [celda.text.strip() for celda in tabla.rows[-1].cells]
        if not "".join(textos_ultima_fila) and tabla.rows[-1] not in filas_rayas_a_borrar:
            filas_rayas_a_borrar.append(tabla.rows[-1])

    for fila in filas_rayas_a_borrar:
        eliminar_fila(fila)


# ==========================================
# GENERADOR: DIRECTOR DE REUNIÓN (RUTA)
# ==========================================

def generar_acta_dr(datos_brutos):
    ruta_base = os.path.dirname(__file__)
    ruta_plantilla = os.path.join(ruta_base, "DR_RUTA_Plantilla_Maestra_Etiquetas.docx")
    doc = docx.Document(ruta_plantilla)

    datos = {}
    for clave, valor in datos_brutos.items():
        if clave.startswith("{{") and clave.endswith("}}"):
            datos[clave] = valor
        else:
            datos[f"{{{{{clave}}}}}"] = valor

    for parrafo in doc.paragraphs:
        reemplazar_manteniendo_formato_estricto(parrafo, datos)

    for tabla in doc.tables:
        limpiar_tabla_por_etiquetas(tabla, datos, SECCIONES_DR_RUTA)

    nombre_competicion = datos.get("{{COMPETICION}}", "Competicion")
    nombre_docx = f"DR {limpiar_nombre_archivo(nombre_competicion)}.docx"
    doc.save(nombre_docx)
    return nombre_docx


# ==========================================
# GENERADOR: JUEZ JEFE DE TRANSPONDEDOR
# ==========================================

def generar_acta_jjt(datos_brutos):
    ruta_base = os.path.dirname(__file__)
    ruta_plantilla = os.path.join(ruta_base, "JJT_PLANTILLA_MAESTRA_DEFINITIVA.docx")
    doc = docx.Document(ruta_plantilla)

    datos_texto = {}
    estado_cuadraditos = {}

    for clave, valor in datos_brutos.items():
        if clave.startswith("CHECK_"):
            estado_cuadraditos[int(clave.replace("CHECK_", ""))] = valor
        else:
            if clave.startswith("{{") and clave.endswith("}}"):
                datos_texto[clave] = valor
            else:
                datos_texto[f"{{{{{clave}}}}}"] = valor

    for parrafo in doc.paragraphs:
        reemplazar_manteniendo_formato_estricto(parrafo, datos_texto)

    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for parrafo in celda.paragraphs:
                    reemplazar_manteniendo_formato_estricto(parrafo, datos_texto)

    for posicion, estado in estado_cuadraditos.items():
        activar_checkbox_por_posicion(doc, posicion, activar=estado)

    nombre_competicion = datos_texto.get("{{COMPETICION}}", "Informe_XXT").strip()
    nombre_docx = f"JJT {limpiar_nombre_archivo(nombre_competicion)}.docx"
    doc.save(nombre_docx)
    return nombre_docx


# ==========================================
# GENERADOR: JURADO TÉCNICO (PISTA)
# ==========================================

FILAS_JURADO_PISTA = {
    2: "DIRECTOR_REUNION",
    3: "DIRECTOR_TECNICO",
    4: "ADJUNTO_D_TECNICA",
    6: "PRESIDENTE_J_APELACION",
    7: "SECRETARIO_J_APELACION",
    8: "VOCAL_J_APELACION",
    10: "JUEZ_ARBITRO",
    11: "AUXILIAR_JUEZ_ARBITRO_1",
    12: "AUXILIAR_JUEZ_ARBITRO_2",
    13: "AUXILIAR_JUEZ_ARBITRO_3",
    14: "AUXILIAR_JUEZ_ARBITRO_4",
    17: "JUEZ_JEFE_CAMARA",
    18: "JUEZ_CAMARA_LLAMADAS_1",
    21: "JUEZ_JEFE_MARCHA",
    22: "JUEZ_DE_MARCHA_1",
    23: "JUEZ_DE_MARCHA_2",
    24: "JUEZ_DE_MARCHA_3",
    25: "JUEZ_DE_MARCHA_4",
    26: "SECRETARIO_DE_MARCHA",
    29: "COORDINADOR_DE_SALIDAS",
    30: "JUEZ_DE_SALIDAS_1",
    31: "JUEZ_DE_SALIDAS_2",
    32: "AYUDANTE_DE_SALIDAS_1",
    35: "JUEZ_JEFE_FOTOFINISH",
    36: "OPERADOR_FOTOFINISH",
    39: "JEFE_CRONOMETRAJE",
    40: "JUEZ_CRONOMETRAJE_1",
    41: "JUEZ_CRONOMETRAJE_2",
    44: "JEFE_LLEGADAS",
    45: "JUEZ_DE_LLEGADAS_1",
    46: "JUEZ_DE_LLEGADAS_2",
    47: "JUEZ_DE_LLEGADAS_3",
    51: "JUEZ_MEDICION_CIENTIFICA",
    54: "OPERADOR_ANEMOMETRO_1",
    55: "OPERADOR_ANEMOMETRO_2",
    58: "JUEZ_JEFE_SALTO_PERTIGA",
    59: "SECRETARIO_SALTO_PERTIGA",
    60: "JUEZ_SALTO_PERTIGA_1",
    63: "JUEZ_JEFE_LANZAMIENTO_MARTILLO",
    64: "SECRETARIO_LANZAMIENTO_MARTILLO",
    65: "JUEZ_CAIDAS_LANZAMIENTO_MARTILLO",
    68: "JUEZ_JEFE_SALTO_LONGITUD",
    69: "SECRETARIO_SALTO_LONGITUD",
    70: "JUEZ_CAIDAS_SALTO_LONGITUD",
    73: "JUEZ_JEFE_SALTO_ALTURA",
    74: "SECRETARIO_SALTO_ALTURA",
    75: "JUEZ_LISTON_SALTO_ALTURA",
    78: "JUEZ_JEFE_LANZAMIENTO_DISCO",
    79: "SECRETARIO_LANZAMIENTO_DISCO",
    80: "JUEZ_CAIDAS_LANZAMIENTO_DISCO",
    83: "JUEZ_JEFE_TRIPLE_SALTO",
    84: "SECRETARIO_TRIPLE_SALTO",
    85: "JUEZ_CAIDAS_TRIPLE_SALTO",
    88: "JUEZ_JEFE_LANZAMIENTO_PESO",
    89: "SECRETARIO_LANZAMIENTO_PESO",
    90: "JUEZ_CAIDAS_LANZAMIENTO_PESO",
    93: "JUEZ_JEFE_LANZAMIENTO_JABALINA",
    94: "SECRETARIO_LANZAMIENTO_JABALINA",
    95: "JUEZ_CAIDAS_LANZAMIENTO_JABALINA",
    98: "SECRETARIA_1",
    99: "SECRETARIA_2",
    100: "SECRETARIA_3",
    103: "ENLACE_DE_FOTO_FINISH",
    104: "AUXILIAR_DE_COMPETICION_1",
    105: "AUXILIAR_DE_COMPETICION_2",
}

HORARIOS_JURADO_PISTA = {
    50: "MEDICION_CIENTIFICA_HORARIO",
    53: "ANEMOMETRO_HORARIO",
    57: "SALTO_PERTIGA_HORARIO",
    62: "LANZAMIENTO_MARTILLO_HORARIO",
    67: "SALTO_LONGITUD_HORARIO",
    72: "SALTO_ALTURA_HORARIO",
    77: "LANZAMIENTO_DISCO_HORARIO",
    82: "TRIPLE_SALTO_HORARIO",
    87: "LANZAMIENTO_PESO_HORARIO",
    92: "LANZAMIENTO_JABALINA_HORARIO",
}

BLOQUES_JURADO_PISTA = [
    {"cabecera": 16, "filas": [17, 18]},
    {"cabecera": 20, "filas": [21, 22, 23, 24, 25, 26]},
    {"cabecera": 28, "filas": [29, 30, 31, 32]},
    {"cabecera": 34, "filas": [35, 36]},
    {"cabecera": 38, "filas": [39, 40, 41]},
    {"cabecera": 43, "filas": [44, 45, 46, 47]},
    {"cabecera": 97, "filas": [98, 99, 100]},
    {"cabecera": 102, "filas": [103, 104, 105]},
]

BLOQUES_CONCURSOS_PISTA = [
    {"cabecera": 50, "filas": [51], "horario": "MEDICION_CIENTIFICA_HORARIO"},
    {"cabecera": 53, "filas": [54, 55], "horario": "ANEMOMETRO_HORARIO"},
    {"cabecera": 57, "filas": [58, 59, 60], "horario": "SALTO_PERTIGA_HORARIO"},
    {"cabecera": 62, "filas": [63, 64, 65], "horario": "LANZAMIENTO_MARTILLO_HORARIO"},
    {"cabecera": 67, "filas": [68, 69, 70], "horario": "SALTO_LONGITUD_HORARIO"},
    {"cabecera": 72, "filas": [73, 74, 75], "horario": "SALTO_ALTURA_HORARIO"},
    {"cabecera": 77, "filas": [78, 79, 80], "horario": "LANZAMIENTO_DISCO_HORARIO"},
    {"cabecera": 82, "filas": [83, 84, 85], "horario": "TRIPLE_SALTO_HORARIO"},
    {"cabecera": 87, "filas": [88, 89, 90], "horario": "LANZAMIENTO_PESO_HORARIO"},
    {"cabecera": 92, "filas": [93, 94, 95], "horario": "LANZAMIENTO_JABALINA_HORARIO"},
]


def rellenar_fila_jurado_pista(fila, prefijo, datos):
    nombre = str(obtener_valor(datos, f"{prefijo}_NOMBRE", "")).strip()
    categoria = str(obtener_valor(datos, f"{prefijo}_CAT", "")).strip()
    panel = str(obtener_valor(datos, f"{prefijo}_PANEL", "")).strip()
    delegacion = str(obtener_valor(datos, f"{prefijo}_DEL", "")).strip()
    marca_4h = normalizar_marca(obtener_valor(datos, f"{prefijo}_4H", ""), estilo="x")
    marca_45h = normalizar_marca(obtener_valor(datos, f"{prefijo}_45H", ""), estilo="si")
    marca_desp = normalizar_marca(obtener_valor(datos, f"{prefijo}_DESP", ""), estilo="si")

    if not any([nombre, categoria, panel, delegacion, marca_4h, marca_45h, marca_desp]):
        return False

    establecer_texto_celda(fila.cells[1], nombre)
    establecer_texto_celda(fila.cells[2], categoria)
    establecer_texto_celda(fila.cells[3], panel)
    establecer_texto_celda(fila.cells[4], delegacion)
    establecer_texto_celda(fila.cells[5], marca_4h)
    establecer_texto_celda(fila.cells[6], marca_45h)
    if len(fila.cells) > 7:
        establecer_texto_celda(fila.cells[7], marca_desp)

    return True


def limpiar_tabla_jurado_pista(tabla, datos):
    filas_originales = list(tabla.rows)
    indices_a_borrar = set()
    filas_con_datos = set()

    for indice, prefijo in FILAS_JURADO_PISTA.items():
        if indice >= len(filas_originales):
            continue
        if rellenar_fila_jurado_pista(filas_originales[indice], prefijo, datos):
            filas_con_datos.add(indice)
        else:
            indices_a_borrar.add(indice)

    for indice, clave_horario in HORARIOS_JURADO_PISTA.items():
        if indice >= len(filas_originales):
            continue
        horario = str(obtener_valor(datos, clave_horario, "")).strip()
        if horario:
            establecer_texto_celda(filas_originales[indice].cells[1], horario)
            filas_con_datos.add(indice)

    for bloque in BLOQUES_JURADO_PISTA:
        tiene_datos_bloque = any(indice in filas_con_datos for indice in bloque["filas"])
        if tiene_datos_bloque:
            filas_con_datos.add(bloque["cabecera"])
        else:
            indices_a_borrar.add(bloque["cabecera"])

    concursos_con_datos = False
    for bloque in BLOQUES_CONCURSOS_PISTA:
        tiene_datos_bloque = bloque["cabecera"] in filas_con_datos or any(
            indice in filas_con_datos for indice in bloque["filas"]
        )
        if tiene_datos_bloque:
            concursos_con_datos = True
            filas_con_datos.add(bloque["cabecera"])
        else:
            indices_a_borrar.add(bloque["cabecera"])

    if concursos_con_datos:
        filas_con_datos.add(49)
    else:
        indices_a_borrar.add(49)

    for indice in sorted(indices_a_borrar, reverse=True):
        if indice < len(filas_originales):
            eliminar_fila(filas_originales[indice])

    filas_vacias_consecutivas = 0
    filas_vacias_a_borrar = []
    for fila in tabla.rows:
        textos_celdas = [celda.text.strip() for celda in fila.cells]
        if not "".join(textos_celdas):
            filas_vacias_consecutivas += 1
            if filas_vacias_consecutivas > 1:
                filas_vacias_a_borrar.append(fila)
        else:
            filas_vacias_consecutivas = 0

    if tabla.rows:
        textos_ultima_fila = [celda.text.strip() for celda in tabla.rows[-1].cells]
        if not "".join(textos_ultima_fila) and tabla.rows[-1] not in filas_vacias_a_borrar:
            filas_vacias_a_borrar.append(tabla.rows[-1])

    for fila in filas_vacias_a_borrar:
        eliminar_fila(fila)


def normalizar_parrafos_finales(datos):
    lineas = []

    bruto = obtener_valor(datos, "PARRAFOS_FINALES", [])
    if isinstance(bruto, str):
        lineas.extend(line.strip() for line in bruto.splitlines() if line.strip())
    elif isinstance(bruto, (list, tuple)):
        lineas.extend(str(linea).strip() for linea in bruto if str(linea).strip())

    for clave_extra in ("OBSERVACIONES_FINALES", "OBSERVACIONES_ADICIONALES"):
        extra = str(obtener_valor(datos, clave_extra, "")).strip()
        if extra:
            lineas.extend(line.strip() for line in extra.splitlines() if line.strip())

    hora_inicio = str(obtener_valor(datos, "HORA_INICIO_COMPETICION", "")).strip()
    hora_fin = str(obtener_valor(datos, "HORA_FIN_COMPETICION", "")).strip()

    if hora_inicio and not any("hora de inicio" in linea.lower() for linea in lineas):
        lineas.append(f"Hora de inicio de la competición: {hora_inicio}")
    if hora_fin and not any("hora de finaliz" in linea.lower() for linea in lineas):
        lineas.append(f"Hora de finalización de la competición: {hora_fin}")

    return lineas


def eliminar_leyenda_jurado_pista(doc):
    for tabla in list(doc.tables):
        texto_tabla = " ".join(celda.text for fila in tabla.rows for celda in fila.cells)
        if "Categorías" in texto_tabla and "Delegación" in texto_tabla and "Paneles" in texto_tabla:
            eliminar_tabla(tabla)
            break


def rellenar_parrafos_finales(doc, lineas):
    if not lineas:
        return

    cola_parrafos_vacios = []
    for parrafo in reversed(doc.paragraphs):
        if parrafo.text.strip():
            if cola_parrafos_vacios:
                break
            continue
        cola_parrafos_vacios.append(parrafo)

    cola_parrafos_vacios.reverse()

    if not cola_parrafos_vacios:
        cola_parrafos_vacios = [doc.add_paragraph("")]

    estilo_referencia = cola_parrafos_vacios[0].style

    for indice, linea in enumerate(lineas):
        if indice < len(cola_parrafos_vacios):
            parrafo_destino = cola_parrafos_vacios[indice]
        else:
            parrafo_destino = doc.add_paragraph("")
            parrafo_destino.style = estilo_referencia
        establecer_texto_manteniendo_formato(parrafo_destino, linea)

    for parrafo_sobrante in cola_parrafos_vacios[len(lineas):]:
        establecer_texto_manteniendo_formato(parrafo_sobrante, "")


def generar_acta_jurado_pista(datos_brutos):
    ruta_base = os.path.dirname(__file__)
    ruta_plantilla = os.path.join(ruta_base, "ModeloJurados_2023.docx")
    doc = docx.Document(ruta_plantilla)

    competicion = str(obtener_valor(datos_brutos, "COMPETICION", "")).strip()
    lugar = str(obtener_valor(datos_brutos, "LUGAR", "")).strip()
    delegacion = str(obtener_valor(datos_brutos, "DELEGACION", "")).strip()
    fecha = str(obtener_valor(datos_brutos, "FECHA", "")).strip()
    jornada = str(obtener_valor(datos_brutos, "JORNADA", "")).strip()

    for parrafo in doc.paragraphs:
        texto = parrafo.text.strip()
        if texto.startswith("Competición:"):
            establecer_texto_manteniendo_formato(parrafo, f"Competición: {competicion}" if competicion else "Competición:")
        elif texto.startswith("Lugar:"):
            establecer_texto_manteniendo_formato(parrafo, f"Lugar: {lugar}" if lugar else "Lugar:")

    tabla_cabecera = doc.tables[0]
    establecer_texto_celda(
        tabla_cabecera.rows[1].cells[1],
        f"Delegación de {delegacion}" if delegacion else "Delegación de",
    )
    establecer_texto_celda(
        tabla_cabecera.rows[1].cells[3],
        f"Fecha: {fecha}" if fecha else "Fecha:",
    )
    establecer_texto_celda(
        tabla_cabecera.rows[2].cells[3],
        f"Jornada: {jornada}" if jornada else "Jornada:",
    )

    limpiar_tabla_jurado_pista(doc.tables[1], datos_brutos)
    eliminar_leyenda_jurado_pista(doc)
    rellenar_parrafos_finales(doc, normalizar_parrafos_finales(datos_brutos))

    nombre_competicion = competicion or "Competicion"
    nombre_docx = f"DR {limpiar_nombre_archivo(nombre_competicion)}.docx"
    doc.save(nombre_docx)
    return nombre_docx


# ==========================================
# CONFIGURACIÓN DE INFORMES
# ==========================================

TIPOS_INFORME = {
    "Director de Reunión (Ruta)": {
        "gem_url": "https://gemini.google.com/gem/1evM9tdpvflf129mCAu7J6TBViiDE7PaL?usp=sharing",
        "generator": generar_acta_dr,
    },
    "Juez Jefe de Transpondedor (JJT)": {
        "gem_url": "https://gemini.google.com/gem/196pc9YorHovWvxJRP3VHzdvL5kuMkHV1?usp=sharing",
        "generator": generar_acta_jjt,
    },
    "Jurado Técnico (Pista)": {
        "gem_url": None,
        "prompt_file": "gem_JURADO_PISTA.txt",
        "generator": generar_acta_jurado_pista,
    },
}


# ==========================================
# INTERFAZ WEB DE STREAMLIT
# ==========================================

st.set_page_config(page_title="Generador de Actas FGA", page_icon="📝")

st.title("Generador de Actas FGA 📝")
st.write("Sigue los pasos para generar el documento oficial en Word:")

tipo_informe = st.selectbox(
    "1️⃣ Selecciona el tipo de informe que quieres generar:",
    ["Elige el tipo de informe", *TIPOS_INFORME.keys()],
)

if tipo_informe != "Elige el tipo de informe":
    config = TIPOS_INFORME[tipo_informe]

    st.write("2️⃣ Obtén los datos con la IA:")
    if config.get("gem_url"):
        st.link_button("Abrir la IA para obtener los datos", config["gem_url"])
        st.caption(
            "ℹ️ Una vez en la IA, escríbele 'hola' o pega el mensaje donde indican que fuiste confirmado para esa carrera."
        )
    else:
        st.info(
            "Para este informe todavía no hay una URL de gem publicada en la app. Puedes usar el prompt base guardado en el repositorio."
        )

        ruta_prompt = os.path.join(os.path.dirname(__file__), config["prompt_file"])
        if os.path.exists(ruta_prompt):
            with open(ruta_prompt, "r", encoding="utf-8") as file:
                contenido_prompt = file.read()

            st.download_button(
                label="Descargar prompt base del gem",
                data=contenido_prompt,
                file_name=config["prompt_file"],
                mime="text/plain",
            )
            with st.expander("Ver prompt base del gem"):
                st.code(contenido_prompt, language="text")
        else:
            st.warning("No encuentro el archivo del prompt base en el servidor.")

    st.write("3️⃣ Pega debajo el texto del diccionario que te ha dado la Inteligencia Artificial.")
    texto_pegado = st.text_area("Pega aquí los datos (Diccionario):", height=320)

    if st.button("4️⃣ Generar Informe"):
        if not texto_pegado.strip():
            st.warning("El cuadro de texto está vacío. Pega los datos primero.")
        else:
            with st.spinner(f"Generando informe de {tipo_informe}..."):
                try:
                    texto_limpio = texto_pegado.replace("\xa0", " ")
                    inicio = texto_limpio.find("{")
                    fin = texto_limpio.rfind("}") + 1

                    if inicio == -1 or fin == 0:
                        st.error("No he encontrado ningún diccionario en el texto. Debe empezar por '{' y acabar por '}'.")
                    else:
                        texto_diccionario = texto_limpio[inicio:fin]
                        datos_procesados = ast.literal_eval(texto_diccionario)
                        archivo_generado = config["generator"](datos_procesados)

                        st.success(f"{tipo_informe} generado con éxito.")
                        with open(archivo_generado, "rb") as file:
                            st.download_button(
                                label="📥 Descargar Documento en Word",
                                data=file,
                                file_name=archivo_generado,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            )

                except SyntaxError:
                    st.error(
                        "Error de formato: el texto pegado tiene un problema de sintaxis (comas, comillas o corchetes)."
                    )
                except FileNotFoundError:
                    st.error(
                        "Error: no se encuentra la plantilla en el servidor. Comprueba que el archivo .docx está subido a GitHub con el nombre correcto."
                    )
                except Exception as e:
                    st.error(f"Error inesperado: {e}")
